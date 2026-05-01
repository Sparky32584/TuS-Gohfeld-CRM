"""
TuS Gohfeld e.V. — Mitgliederverwaltung
Flask Backend (app.py)

Autor:       Stephan Schwenn  (GitHub: Sparky32584)
Repository:  Sparky32584/TuS-Gohfeld-CRM  (privat)
Lizenz:      intern, nicht für Weitergabe bestimmt

Lokale, DSGVO-konforme Mitgliederverwaltung. Alle Daten bleiben auf dem Gerät,
keine externen Server, kein Tracking.

Start:
    python3 app.py      →  Flask Dev Server auf http://127.0.0.1:5000
    python3 run.py      →  Native pywebview-Fenster

Umgebungsvariable:
    TUS_DATA_DIR        →  Pfad für mitglieder.db, dokumente/, settings.json
                           Standard: ./data  (relativ zu app.py)
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
import sqlite3
import tempfile
import time
import xml.etree.ElementTree as ET
from calendar import monthrange
from contextlib import contextmanager
from datetime import datetime, date
from pathlib import Path
from typing import Any, Iterable

from flask import (
    Flask, abort, g, jsonify, render_template, request,
    send_file, send_from_directory,
)
from werkzeug.utils import secure_filename

# ─── Optionale Abhängigkeiten (App startet auch ohne OCR-/PDF-Stack) ────
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image, ImageEnhance, ImageFilter
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import docx as python_docx  # python-docx
except ImportError:
    python_docx = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm, cm
    from reportlab.lib.colors import HexColor, black, white, lightgrey
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# ═════════════════════════════════════════════════════════════════════════
#  KONFIGURATION
# ═════════════════════════════════════════════════════════════════════════

APP_ROOT = Path(__file__).resolve().parent
DATA_DIR = Path(os.environ.get("TUS_DATA_DIR", APP_ROOT / "data")).expanduser().resolve()
DB_PATH  = DATA_DIR / "mitglieder.db"
DOC_DIR  = DATA_DIR / "dokumente"
SETTINGS = DATA_DIR / "settings.json"

ALLOWED_EXT = {"pdf", "png", "jpg", "jpeg", "doc", "docx", "odt", "txt", "csv", "xml"}
MAX_SIZE    = 100 * 1024 * 1024   # 100 MB

# Verzeichnisstruktur initialisieren
for p in (DATA_DIR, DOC_DIR):
    p.mkdir(parents=True, exist_ok=True)

app = Flask(__name__, template_folder="templates", static_folder="static")
app.config["MAX_CONTENT_LENGTH"] = MAX_SIZE
app.config["JSON_AS_ASCII"]      = False

# ═════════════════════════════════════════════════════════════════════════
#  STANDARD-EINSTELLUNGEN  (werden in settings.json persistiert)
# ═════════════════════════════════════════════════════════════════════════

DEFAULT_SETTINGS: dict[str, Any] = {
    "colors": {
        "primary":    "#3156a3",
        "sidebar_bg": "#1e2a4a",
        "page_bg":    "#f5f6fa",
        "card_bg":    "#ffffff",
        "text":       "#1a1f36",
        "border":     "#e2e5ef",
    },
    "display": {
        "show_member_number": True,
        "show_age":           True,
        "show_gender":        True,
        "show_email":         False,
        "show_phone":         False,
        "stale_detection":    True,
    },
    "doc_types": [
        "Datenschutzerklärung",
        "SEPA-Lastschriftmandat",
        "Mitgliedsantrag",
        "Einverständniserklärung",
        "Ärztliches Attest",
        "Sonstiges",
    ],
    "color_meanings": {
        "red":    "Beitrag ausstehend",
        "orange": "Daten unvollständig",
        "yellow": "Kontaktversuch offen",
        "green":  "Alles vollständig",
        "blue":   "Vorstand / Funktionsträger",
        "purple": "Ehrenmitglied",
    },
}

VALID_COLORS = {"red", "orange", "yellow", "green", "blue", "purple", ""}
APP_STATUS   = {"Nicht installiert", "Installiert", "Aktiv"}


# ═════════════════════════════════════════════════════════════════════════
#  DATENBANK — Verbindung + Schema-Migration
# ═════════════════════════════════════════════════════════════════════════

def db_conn() -> sqlite3.Connection:
    """Liefert eine SQLite-Verbindung mit WAL-Modus + Foreign Keys."""
    conn = sqlite3.connect(DB_PATH, timeout=20.0, isolation_level=None)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode = WAL;")
    conn.execute("PRAGMA synchronous  = NORMAL;")
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn


@contextmanager
def db() -> Iterable[sqlite3.Connection]:
    """Kontext-Manager: commit bei Erfolg, rollback bei Fehler."""
    conn = db_conn()
    try:
        yield conn
    except Exception:
        try:
            conn.execute("ROLLBACK")
        except sqlite3.Error:
            pass
        raise
    finally:
        conn.close()


def _table_columns(conn: sqlite3.Connection, table: str) -> set[str]:
    return {row["name"] for row in conn.execute(f"PRAGMA table_info({table})")}


def _ensure_column(conn: sqlite3.Connection, table: str, column: str, ddl: str) -> None:
    """Fügt eine Spalte hinzu, wenn sie noch nicht existiert (Migrations-Helper)."""
    if column not in _table_columns(conn, table):
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {ddl}")


def init_db() -> None:
    """Legt Tabellen/Indexe an und ergänzt fehlende Spalten bei bestehenden DBs."""
    with db() as c:
        # ── members ─────────────────────────────────────────────────────
        c.executescript("""
            CREATE TABLE IF NOT EXISTS members (
                id             INTEGER PRIMARY KEY AUTOINCREMENT,
                member_number  TEXT UNIQUE,
                first_name     TEXT NOT NULL,
                last_name      TEXT NOT NULL,
                birth_date     TEXT,
                gender         TEXT DEFAULT '',
                email          TEXT DEFAULT '',
                phone          TEXT DEFAULT '',
                street         TEXT DEFAULT '',
                city           TEXT DEFAULT 'Löhne',
                zip_code       TEXT DEFAULT '32584',
                entry_date     TEXT,
                exit_date      TEXT,
                color_mark     TEXT DEFAULT '',
                app_status     TEXT DEFAULT 'Nicht installiert',
                notes          TEXT DEFAULT '',
                is_active      INTEGER DEFAULT 1,
                created_at     TEXT,
                updated_at     TEXT
            );
            CREATE INDEX IF NOT EXISTS idx_members_name
                ON members (last_name, first_name);
            CREATE INDEX IF NOT EXISTS idx_members_active
                ON members (is_active);
        """)

        # Migrationen für bestehende DBs (fehlertolerant)
        for col, ddl in [
            ("gender",      "TEXT DEFAULT ''"),
            ("color_mark",  "TEXT DEFAULT ''"),
            ("app_status",  "TEXT DEFAULT 'Nicht installiert'"),
            ("notes",       "TEXT DEFAULT ''"),
            ("exit_date",   "TEXT"),
            ("is_active",   "INTEGER DEFAULT 1"),
            ("created_at",  "TEXT"),
            ("updated_at",  "TEXT"),
        ]:
            _ensure_column(c, "members", col, ddl)

        # ── groups_tbl ──────────────────────────────────────────────────
        c.executescript("""
            CREATE TABLE IF NOT EXISTS groups_tbl (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                name        TEXT NOT NULL,
                description TEXT DEFAULT '',
                color       TEXT DEFAULT '#4CAF50',
                level       INTEGER DEFAULT 1,
                parent_id   INTEGER DEFAULT NULL
                             REFERENCES groups_tbl(id) ON DELETE CASCADE,
                sort_order  INTEGER DEFAULT 0
            );
            CREATE INDEX IF NOT EXISTS idx_groups_parent
                ON groups_tbl (parent_id);
        """)
        for col, ddl in [
            ("description", "TEXT DEFAULT ''"),
            ("color",       "TEXT DEFAULT '#4CAF50'"),
            ("level",       "INTEGER DEFAULT 1"),
            ("parent_id",   "INTEGER"),
            ("sort_order",  "INTEGER DEFAULT 0"),
        ]:
            _ensure_column(c, "groups_tbl", col, ddl)

        # ── member_groups ───────────────────────────────────────────────
        c.execute("""
            CREATE TABLE IF NOT EXISTS member_groups (
                member_id INTEGER NOT NULL
                            REFERENCES members(id)    ON DELETE CASCADE,
                group_id  INTEGER NOT NULL
                            REFERENCES groups_tbl(id) ON DELETE CASCADE,
                PRIMARY KEY (member_id, group_id)
            );
        """)

        # ── documents ───────────────────────────────────────────────────
        c.executescript("""
            CREATE TABLE IF NOT EXISTS documents (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                member_id     INTEGER NOT NULL
                                REFERENCES members(id) ON DELETE CASCADE,
                doc_type      TEXT NOT NULL DEFAULT 'Sonstiges',
                original_name TEXT NOT NULL,
                stored_name   TEXT NOT NULL,
                uploaded_at   TEXT
            );
            CREATE INDEX IF NOT EXISTS idx_documents_member
                ON documents (member_id);
        """)


# ═════════════════════════════════════════════════════════════════════════
#  EINSTELLUNGEN — JSON-Datei in DATA_DIR
# ═════════════════════════════════════════════════════════════════════════

def _deep_merge(base: dict, override: dict) -> dict:
    """Rekursives Merge — fehlende Keys aus `base` werden ergänzt."""
    out = dict(base)
    for k, v in override.items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _deep_merge(out[k], v)
        else:
            out[k] = v
    return out


def load_settings() -> dict:
    """Lädt settings.json, füllt fehlende Keys mit Defaults auf."""
    if not SETTINGS.exists():
        save_settings(DEFAULT_SETTINGS)
        return dict(DEFAULT_SETTINGS)
    try:
        user = json.loads(SETTINGS.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return dict(DEFAULT_SETTINGS)
    # Defaults bilden Basis, User-Werte überschreiben
    return _deep_merge(DEFAULT_SETTINGS, user)


def save_settings(data: dict) -> None:
    SETTINGS.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


# ═════════════════════════════════════════════════════════════════════════
#  NAMENS-HEURISTIK  (Geschlechts-Vermutung anhand Vorname)
# ═════════════════════════════════════════════════════════════════════════

MALE_NAMES = {
    # Deutsch
    "alexander", "andreas", "alfred", "alfons", "anton", "arthur", "axel",
    "benjamin", "bernd", "bernhard", "björn", "bodo", "bruno", "burkhard",
    "christian", "christoph", "claus", "daniel", "david", "dennis", "detlef",
    "dieter", "dirk", "dominik", "eberhard", "edgar", "egon", "elias",
    "emil", "erich", "erik", "ernst", "erwin", "eugen", "felix", "florian",
    "frank", "franz", "friedrich", "fritz", "gabriel", "georg", "gerd",
    "gerhard", "günter", "gustav", "hannes", "hans", "harald", "hartmut",
    "hartwig", "heinrich", "heinz", "helge", "helmut", "henning", "henrik",
    "herbert", "hermann", "holger", "horst", "hubert", "ingo", "jens",
    "joachim", "jochen", "johannes", "jonas", "jörg", "josef", "jürgen",
    "karl", "karsten", "kevin", "klaus", "knut", "konrad", "kurt", "lars",
    "leo", "leon", "lothar", "lucas", "lukas", "ludwig", "manfred", "marc",
    "marco", "marcel", "mario", "mark", "markus", "martin", "matthias",
    "max", "maximilian", "michael", "moritz", "nico", "niklas", "nils",
    "norbert", "olaf", "oliver", "oskar", "otto", "patrick", "paul", "peter",
    "philipp", "ralf", "reiner", "reinhard", "reinhold", "richard", "robert",
    "roland", "rolf", "roman", "rudolf", "rüdiger", "sebastian", "siegfried",
    "simon", "stefan", "stephan", "sven", "theo", "thomas", "tim", "timo",
    "tobias", "tom", "torsten", "udo", "ulrich", "uwe", "viktor", "volker",
    "walter", "werner", "wilhelm", "willi", "willy", "wolfgang", "yannick",
    # International
    "adrian", "alex", "andrea", "antonio", "carlos", "chris", "james",
    "john", "luis", "michel", "mustafa", "nicolas", "pablo", "sergej",
}

FEMALE_NAMES = {
    "agnes", "alexandra", "alina", "amelie", "andrea", "angela", "angelika",
    "anja", "anke", "anna", "anne", "annegret", "annette", "annika", "antje",
    "astrid", "barbara", "beate", "bettina", "birgit", "brigitte", "britta",
    "carina", "carmen", "carolin", "caroline", "christa", "christel",
    "christiane", "christine", "claudia", "cornelia", "dagmar", "daniela",
    "diana", "dorothea", "edith", "elfriede", "elisabeth", "elke", "ellen",
    "elsa", "elvira", "emilia", "emma", "erika", "eva", "franziska",
    "frauke", "frieda", "friederike", "gabi", "gabriele", "gerda", "gerlinde",
    "gertrud", "gisela", "greta", "gudrun", "hannah", "hanna", "hannelore",
    "heidi", "heike", "helene", "helga", "henriette", "hilde", "hildegard",
    "ilona", "ilse", "inge", "ingeborg", "ingrid", "irene", "iris", "irmgard",
    "isabel", "jana", "janine", "jasmin", "jenny", "jennifer", "jessica",
    "johanna", "julia", "juliane", "jutta", "karin", "karla", "karolin",
    "katharina", "kathrin", "katja", "katrin", "kerstin", "kirsten", "klara",
    "larissa", "laura", "lea", "leonie", "lena", "lieselotte", "lina",
    "linda", "lisa", "lotte", "luise", "madeleine", "maike", "manuela",
    "mareike", "margit", "margot", "margarete", "maria", "marianne",
    "marie", "marina", "marion", "marlene", "martha", "martina", "mathilda",
    "mechthild", "melanie", "meike", "mia", "michaela", "mila", "monika",
    "nadine", "nadja", "natalie", "nicole", "nina", "olga", "paula",
    "petra", "pia", "regina", "renate", "rita", "romy", "rosa", "rosemarie",
    "ruth", "sabine", "sabrina", "sandra", "sara", "sarah", "saskia",
    "silke", "silvia", "simone", "sonja", "sophia", "sophie", "stefanie",
    "stephanie", "steffi", "susanne", "svenja", "sylvia", "tabea", "tamara",
    "tanja", "theresa", "tina", "ulla", "ulrike", "ursula", "ute", "vanessa",
    "vera", "veronika", "viktoria", "waltraud", "wiebke", "yvonne",
    # International
    "alice", "ana", "aysegül", "clara", "fatima", "julie", "lara",
    "lucia", "maria", "natalia", "sabina", "sofia",
}


def guess_gender(first_name: str) -> str:
    """Vermutet das Geschlecht anhand des Vornamens.

    Rückgabe: 'männlich', 'weiblich' oder '' (unbekannt).
    """
    if not first_name:
        return ""
    # Erster Vorname (bei "Hans Peter" nur "hans")
    name = first_name.strip().split()[0].lower()

    if name in MALE_NAMES:
        return "männlich"
    if name in FEMALE_NAMES:
        return "weiblich"

    # Endungs-Heuristik als Fallback
    female_endings = ("a", "e", "ine", "ette", "ina", "elle", "ana", "iya")
    male_endings   = ("er", "us", "ard", "helm", "bert", "fried", "rich", "old")
    if any(name.endswith(e) for e in male_endings):
        return "männlich"
    if any(name.endswith(e) for e in female_endings):
        return "weiblich"

    return ""


# ═════════════════════════════════════════════════════════════════════════
#  MITGLIEDSNUMMER — Auto-Generierung
# ═════════════════════════════════════════════════════════════════════════

def next_member_number(conn: sqlite3.Connection) -> str:
    """Gibt die nächste 5-stellige Mitgliedsnummer zurück (z.B. '00042').

    10-stellige Nummern aus TuS-Importen werden ignoriert, damit interne
    Nummern nicht in den 10-stelligen Bereich springen.
    """
    row = conn.execute("""
        SELECT MAX(CAST(member_number AS INTEGER)) AS max_num
        FROM members
        WHERE member_number GLOB '[0-9][0-9][0-9][0-9][0-9]'
    """).fetchone()
    nxt = (row["max_num"] or 0) + 1
    return f"{nxt:05d}"


# ═════════════════════════════════════════════════════════════════════════
#  KARTEILEICHEN-ERKENNUNG
# ═════════════════════════════════════════════════════════════════════════

def is_karteileiche(m: sqlite3.Row | dict, group_count: int) -> bool:
    """Ein Mitglied gilt als Karteileiche, wenn:
    - Aktiv UND Eintritt > 5 Jahre her UND keiner Gruppe zugeordnet
    - ODER: Aktiv UND kein Eintrittsdatum UND keine E-Mail UND kein Telefon
    """
    if not (m["is_active"] if isinstance(m, sqlite3.Row) else m.get("is_active")):
        return False

    entry = m["entry_date"] if isinstance(m, sqlite3.Row) else m.get("entry_date")
    email = (m["email"]     if isinstance(m, sqlite3.Row) else m.get("email", "")) or ""
    phone = (m["phone"]     if isinstance(m, sqlite3.Row) else m.get("phone", "")) or ""

    if entry:
        try:
            years = (date.today() - date.fromisoformat(entry)).days / 365.25
            if years > 5 and group_count == 0:
                return True
        except ValueError:
            pass
    else:
        if not email.strip() and not phone.strip():
            return True

    return False


# ═════════════════════════════════════════════════════════════════════════
#  HILFSFUNKTIONEN
# ═════════════════════════════════════════════════════════════════════════

def now_local() -> str:
    """Aktueller lokaler Zeitstempel im SQLite-kompatiblen Format."""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def parse_date(value: str | None) -> str | None:
    """Normalisiert verschiedene Datumsformate auf YYYY-MM-DD.

    Akzeptiert: DD.MM.YYYY, YYYY-MM-DD, DD/MM/YYYY, DD-MM-YYYY,
    YYYYMMDD, DD.MM.YY, MM/DD/YYYY.
    Rückgabe None, wenn leer oder nicht parsebar.
    """
    if not value:
        return None
    s = str(value).strip()
    if not s:
        return None

    formats = (
        "%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y", "%d-%m-%Y",
        "%Y%m%d",  "%d.%m.%y", "%m/%d/%Y",
    )
    for fmt in formats:
        try:
            return datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            continue
    return None


def calc_age(birth: str | None) -> int | None:
    if not birth:
        return None
    try:
        b = date.fromisoformat(birth)
    except ValueError:
        return None
    today = date.today()
    return today.year - b.year - ((today.month, today.day) < (b.month, b.day))


def _member_row_to_dict(m: sqlite3.Row, groups: list[dict] | None = None,
                       docs: list[dict] | None = None) -> dict:
    d = dict(m)
    d["age"]      = calc_age(d.get("birth_date"))
    if groups is not None:
        d["groups"] = groups
    if docs is not None:
        d["documents"] = docs
    return d


def _load_member_groups(conn: sqlite3.Connection, member_id: int) -> list[dict]:
    rows = conn.execute("""
        SELECT g.id, g.name, g.level, g.parent_id, g.color
        FROM member_groups mg
        JOIN groups_tbl g ON g.id = mg.group_id
        WHERE mg.member_id = ?
        ORDER BY g.level, g.name
    """, (member_id,)).fetchall()
    return [dict(r) for r in rows]


def _load_member_documents(conn: sqlite3.Connection, member_id: int) -> list[dict]:
    rows = conn.execute("""
        SELECT id, doc_type, original_name, stored_name, uploaded_at
        FROM documents
        WHERE member_id = ?
        ORDER BY uploaded_at DESC
    """, (member_id,)).fetchall()
    return [dict(r) for r in rows]


# ═════════════════════════════════════════════════════════════════════════
#  FEHLER-HANDLER
# ═════════════════════════════════════════════════════════════════════════

@app.errorhandler(400)
def _err_400(e): return jsonify(error=str(e.description)), 400

@app.errorhandler(404)
def _err_404(e): return jsonify(error=str(e.description)), 404

@app.errorhandler(409)
def _err_409(e): return jsonify(error=str(e.description)), 409

@app.errorhandler(413)
def _err_413(e): return jsonify(error="Datei zu groß (max. 100 MB)"), 413

@app.errorhandler(500)
def _err_500(e):
    app.logger.exception("Interner Serverfehler")
    return jsonify(error="Interner Serverfehler"), 500


# ═════════════════════════════════════════════════════════════════════════
#  ROUTE:  Frontend  (Single-Page-App)
# ═════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/favicon.ico")
def favicon():
    fav = APP_ROOT / "static" / "logo.png"
    if fav.exists():
        return send_from_directory(fav.parent, fav.name)
    abort(404)


# ═════════════════════════════════════════════════════════════════════════
#  API:  MITGLIEDER
# ═════════════════════════════════════════════════════════════════════════

@app.route("/api/members", methods=["GET"])
def api_members_list():
    search       = (request.args.get("search") or "").strip()
    group_id     = request.args.get("group_id", type=int)
    color_mark   = request.args.get("color_mark")
    app_status   = request.args.get("app_status")
    active_only  = request.args.get("active_only", "true").lower() == "true"

    sql = "SELECT m.* FROM members m"
    joins, where, params = [], [], []

    if group_id:
        joins.append("JOIN member_groups mg ON mg.member_id = m.id")
        where.append("mg.group_id = ?")
        params.append(group_id)

    if active_only:
        where.append("m.is_active = 1")

    if search:
        where.append("""(
            m.first_name    LIKE ? OR m.last_name LIKE ? OR
            m.member_number LIKE ? OR m.email     LIKE ? OR
            m.phone         LIKE ? OR m.city      LIKE ?
        )""")
        like = f"%{search}%"
        params.extend([like] * 6)

    if color_mark is not None and color_mark in VALID_COLORS:
        where.append("m.color_mark = ?")
        params.append(color_mark)

    if app_status in APP_STATUS:
        where.append("m.app_status = ?")
        params.append(app_status)

    if joins:
        sql += " " + " ".join(joins)
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY m.last_name COLLATE NOCASE, m.first_name COLLATE NOCASE"

    with db() as c:
        rows = c.execute(sql, params).fetchall()
        result = []
        for r in rows:
            groups = _load_member_groups(c, r["id"])
            d = _member_row_to_dict(r, groups=groups)
            d["karteileiche"] = is_karteileiche(r, len(groups))
            result.append(d)

    return jsonify(result)


@app.route("/api/members", methods=["POST"])
def api_member_create():
    data = request.get_json(silent=True) or {}
    first = (data.get("first_name") or "").strip()
    last  = (data.get("last_name")  or "").strip()
    if not first or not last:
        abort(400, "Vor- und Nachname sind Pflichtfelder")

    with db() as c:
        c.execute("BEGIN")
        number = (data.get("member_number") or "").strip()
        if not number:
            number = next_member_number(c)

        # Duplikat-Check
        if c.execute("SELECT 1 FROM members WHERE member_number = ?",
                     (number,)).fetchone():
            abort(409, f"Mitgliedsnummer {number} existiert bereits")

        gender = data.get("gender") or guess_gender(first)
        now = now_local()

        cur = c.execute("""
            INSERT INTO members (
                member_number, first_name, last_name, birth_date, gender,
                email, phone, street, city, zip_code, entry_date, exit_date,
                color_mark, app_status, notes, is_active, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            number, first, last,
            parse_date(data.get("birth_date")),
            gender,
            (data.get("email") or "").strip(),
            (data.get("phone") or "").strip(),
            (data.get("street") or "").strip(),
            (data.get("city") or "Löhne").strip(),
            (data.get("zip_code") or "32584").strip(),
            parse_date(data.get("entry_date")),
            parse_date(data.get("exit_date")),
            data.get("color_mark") or "",
            data.get("app_status") if data.get("app_status") in APP_STATUS
                                   else "Nicht installiert",
            data.get("notes") or "",
            0 if data.get("exit_date") else 1,
            now, now,
        ))
        member_id = cur.lastrowid

        # Gruppen-Zuordnung
        group_ids = data.get("group_ids") or []
        if isinstance(group_ids, list) and group_ids:
            c.executemany(
                "INSERT OR IGNORE INTO member_groups (member_id, group_id) VALUES (?, ?)",
                [(member_id, int(g)) for g in group_ids if str(g).isdigit()],
            )
        c.execute("COMMIT")

        row = c.execute("SELECT * FROM members WHERE id = ?", (member_id,)).fetchone()
        groups = _load_member_groups(c, member_id)

    return jsonify(_member_row_to_dict(row, groups=groups, docs=[])), 201


@app.route("/api/members/<int:member_id>", methods=["GET"])
def api_member_detail(member_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM members WHERE id = ?", (member_id,)).fetchone()
        if not row:
            abort(404, "Mitglied nicht gefunden")
        groups = _load_member_groups(c, member_id)
        docs   = _load_member_documents(c, member_id)
    return jsonify(_member_row_to_dict(row, groups=groups, docs=docs))


@app.route("/api/members/<int:member_id>", methods=["PUT"])
def api_member_update(member_id: int):
    data = request.get_json(silent=True) or {}
    with db() as c:
        row = c.execute("SELECT * FROM members WHERE id = ?", (member_id,)).fetchone()
        if not row:
            abort(404, "Mitglied nicht gefunden")

        # Aktualisierbare Felder
        fields = [
            "first_name", "last_name", "birth_date", "gender", "email",
            "phone", "street", "city", "zip_code", "entry_date", "exit_date",
            "color_mark", "app_status", "notes", "member_number",
        ]
        updates: dict[str, Any] = {}
        for f in fields:
            if f in data:
                val = data[f]
                if f in ("birth_date", "entry_date", "exit_date"):
                    val = parse_date(val)
                elif f == "color_mark" and val not in VALID_COLORS:
                    val = ""
                elif f == "app_status" and val not in APP_STATUS:
                    val = "Nicht installiert"
                elif isinstance(val, str):
                    val = val.strip()
                updates[f] = val

        # Pflichtfelder
        if "first_name" in updates and not updates["first_name"]:
            abort(400, "Vorname darf nicht leer sein")
        if "last_name"  in updates and not updates["last_name"]:
            abort(400, "Nachname darf nicht leer sein")

        # is_active aus exit_date ableiten
        if "exit_date" in updates:
            updates["is_active"] = 0 if updates["exit_date"] else 1

        # Mitgliedsnummer-Konflikt prüfen
        if "member_number" in updates and updates["member_number"] != row["member_number"]:
            dup = c.execute(
                "SELECT 1 FROM members WHERE member_number = ? AND id <> ?",
                (updates["member_number"], member_id),
            ).fetchone()
            if dup:
                abort(409, f"Mitgliedsnummer {updates['member_number']} existiert bereits")

        updates["updated_at"] = now_local()

        if updates:
            set_clause = ", ".join(f"{k} = ?" for k in updates)
            c.execute("BEGIN")
            c.execute(f"UPDATE members SET {set_clause} WHERE id = ?",
                      (*updates.values(), member_id))
            # Gruppen-Zuordnung (komplett ersetzen)
            if "group_ids" in data and isinstance(data["group_ids"], list):
                c.execute("DELETE FROM member_groups WHERE member_id = ?", (member_id,))
                ids = [int(g) for g in data["group_ids"] if str(g).isdigit()]
                if ids:
                    c.executemany(
                        "INSERT OR IGNORE INTO member_groups (member_id, group_id) VALUES (?, ?)",
                        [(member_id, g) for g in ids],
                    )
            c.execute("COMMIT")
        elif "group_ids" in data and isinstance(data["group_ids"], list):
            c.execute("BEGIN")
            c.execute("DELETE FROM member_groups WHERE member_id = ?", (member_id,))
            ids = [int(g) for g in data["group_ids"] if str(g).isdigit()]
            if ids:
                c.executemany(
                    "INSERT OR IGNORE INTO member_groups (member_id, group_id) VALUES (?, ?)",
                    [(member_id, g) for g in ids],
                )
            c.execute("COMMIT")

        row = c.execute("SELECT * FROM members WHERE id = ?", (member_id,)).fetchone()
        groups = _load_member_groups(c, member_id)
        docs   = _load_member_documents(c, member_id)

    return jsonify(_member_row_to_dict(row, groups=groups, docs=docs))


@app.route("/api/members/<int:member_id>", methods=["DELETE"])
def api_member_delete(member_id: int):
    with db() as c:
        row = c.execute("SELECT 1 FROM members WHERE id = ?", (member_id,)).fetchone()
        if not row:
            abort(404, "Mitglied nicht gefunden")

        # Zugehörige Dokumentdateien löschen
        docs = c.execute(
            "SELECT stored_name FROM documents WHERE member_id = ?", (member_id,)
        ).fetchall()
        for d in docs:
            fp = DOC_DIR / d["stored_name"]
            if fp.exists():
                try:
                    fp.unlink()
                except OSError:
                    app.logger.warning("Konnte Datei nicht löschen: %s", fp)

        c.execute("BEGIN")
        c.execute("DELETE FROM members WHERE id = ?", (member_id,))
        c.execute("COMMIT")
    return jsonify(ok=True)


@app.route("/api/members/<int:member_id>/color", methods=["PATCH"])
def api_member_color(member_id: int):
    data  = request.get_json(silent=True) or {}
    color = (data.get("color_mark") or "").lower()
    if color not in VALID_COLORS:
        abort(400, f"Ungültige Farbe: {color}")

    with db() as c:
        if not c.execute("SELECT 1 FROM members WHERE id = ?", (member_id,)).fetchone():
            abort(404, "Mitglied nicht gefunden")
        c.execute("BEGIN")
        c.execute("UPDATE members SET color_mark = ?, updated_at = ? WHERE id = ?",
                  (color, now_local(), member_id))
        c.execute("COMMIT")
    return jsonify(ok=True, color_mark=color)


@app.route("/api/members/bulk-delete", methods=["POST"])
def api_members_bulk_delete():
    data = request.get_json(silent=True) or {}
    ids  = data.get("ids") or []
    if not isinstance(ids, list) or not ids:
        abort(400, "Keine IDs übergeben")
    ids = [int(i) for i in ids if str(i).isdigit()]
    if not ids:
        abort(400, "Keine gültigen IDs übergeben")

    with db() as c:
        placeholders = ",".join("?" * len(ids))
        # Dateien löschen
        docs = c.execute(
            f"SELECT stored_name FROM documents WHERE member_id IN ({placeholders})", ids
        ).fetchall()
        for d in docs:
            fp = DOC_DIR / d["stored_name"]
            if fp.exists():
                try: fp.unlink()
                except OSError: pass

        c.execute("BEGIN")
        c.execute(f"DELETE FROM members WHERE id IN ({placeholders})", ids)
        c.execute("COMMIT")
    return jsonify(ok=True, deleted=len(ids))


# ═════════════════════════════════════════════════════════════════════════
#  API:  GRUPPEN  (3-stufige Hierarchie)
# ═════════════════════════════════════════════════════════════════════════

@app.route("/api/groups", methods=["GET"])
def api_groups_list():
    with db() as c:
        rows = c.execute("""
            SELECT g.*, (
                SELECT COUNT(*) FROM member_groups mg WHERE mg.group_id = g.id
            ) AS member_count
            FROM groups_tbl g
            ORDER BY g.level, g.sort_order, g.name COLLATE NOCASE
        """).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/groups/tree", methods=["GET"])
def api_groups_tree():
    """Liefert Abteilungen → Sportarten → Gruppen als Baum mit Mitgliederzahlen."""
    with db() as c:
        rows = c.execute("""
            SELECT g.*, (
                SELECT COUNT(*) FROM member_groups mg WHERE mg.group_id = g.id
            ) AS direct_members
            FROM groups_tbl g
            ORDER BY g.level, g.sort_order, g.name COLLATE NOCASE
        """).fetchall()

    by_id: dict[int, dict] = {}
    roots: list[dict] = []
    for r in rows:
        node = dict(r)
        node["children"] = []
        node["total_members"] = node["direct_members"]
        by_id[node["id"]] = node

    # Eltern-Kind verknüpfen
    for node in by_id.values():
        pid = node.get("parent_id")
        if pid and pid in by_id:
            by_id[pid]["children"].append(node)
        else:
            roots.append(node)

    # Rekursiv total_members aufsummieren
    def aggregate(n: dict) -> int:
        total = n["direct_members"]
        for ch in n["children"]:
            total += aggregate(ch)
        n["total_members"] = total
        return total
    for r in roots:
        aggregate(r)

    return jsonify(roots)


@app.route("/api/groups", methods=["POST"])
def api_group_create():
    data = request.get_json(silent=True) or {}
    name = (data.get("name") or "").strip()
    if not name:
        abort(400, "Gruppenname darf nicht leer sein")

    level     = int(data.get("level") or 1)
    level     = max(1, min(3, level))
    parent_id = data.get("parent_id")
    if parent_id in ("", None, 0):
        parent_id = None
    else:
        parent_id = int(parent_id)

    with db() as c:
        # Farbe vom Eltern erben, falls nicht angegeben
        color = data.get("color")
        if parent_id and not color:
            parent = c.execute("SELECT color FROM groups_tbl WHERE id = ?",
                               (parent_id,)).fetchone()
            if parent:
                color = parent["color"]
        if not color:
            color = "#4CAF50"

        # Level aus Elternteil ableiten
        if parent_id:
            parent = c.execute("SELECT level FROM groups_tbl WHERE id = ?",
                               (parent_id,)).fetchone()
            if not parent:
                abort(400, "Eltern-Gruppe nicht gefunden")
            level = min(3, parent["level"] + 1)

        c.execute("BEGIN")
        cur = c.execute("""
            INSERT INTO groups_tbl (name, description, color, level, parent_id, sort_order)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (
            name,
            data.get("description") or "",
            color,
            level,
            parent_id,
            int(data.get("sort_order") or 0),
        ))
        c.execute("COMMIT")
        row = c.execute("SELECT * FROM groups_tbl WHERE id = ?",
                        (cur.lastrowid,)).fetchone()
    return jsonify(dict(row)), 201


@app.route("/api/groups/<int:group_id>", methods=["PUT"])
def api_group_update(group_id: int):
    data = request.get_json(silent=True) or {}
    with db() as c:
        row = c.execute("SELECT * FROM groups_tbl WHERE id = ?",
                        (group_id,)).fetchone()
        if not row:
            abort(404, "Gruppe nicht gefunden")

        updates: dict[str, Any] = {}
        for f in ("name", "description", "color", "sort_order"):
            if f in data:
                val = data[f]
                if f == "sort_order":
                    val = int(val or 0)
                elif isinstance(val, str):
                    val = val.strip()
                if f == "name" and not val:
                    abort(400, "Gruppenname darf nicht leer sein")
                updates[f] = val

        if updates:
            set_clause = ", ".join(f"{k} = ?" for k in updates)
            c.execute("BEGIN")
            c.execute(f"UPDATE groups_tbl SET {set_clause} WHERE id = ?",
                      (*updates.values(), group_id))

            # Farb-Vererbung: wenn Farbe geändert, auch Kinder aktualisieren
            if "color" in updates:
                _cascade_color(c, group_id, updates["color"])

            c.execute("COMMIT")
        row = c.execute("SELECT * FROM groups_tbl WHERE id = ?",
                        (group_id,)).fetchone()
    return jsonify(dict(row))


def _cascade_color(conn: sqlite3.Connection, parent_id: int, color: str) -> None:
    """Übernimmt die Farbe rekursiv in alle Kind- und Enkelgruppen."""
    children = conn.execute(
        "SELECT id FROM groups_tbl WHERE parent_id = ?", (parent_id,)
    ).fetchall()
    for ch in children:
        conn.execute("UPDATE groups_tbl SET color = ? WHERE id = ?",
                     (color, ch["id"]))
        _cascade_color(conn, ch["id"], color)


@app.route("/api/groups/<int:group_id>", methods=["DELETE"])
def api_group_delete(group_id: int):
    with db() as c:
        if not c.execute("SELECT 1 FROM groups_tbl WHERE id = ?",
                         (group_id,)).fetchone():
            abort(404, "Gruppe nicht gefunden")
        c.execute("BEGIN")
        # Cascade via FK-Constraint (ON DELETE CASCADE)
        c.execute("DELETE FROM groups_tbl WHERE id = ?", (group_id,))
        c.execute("COMMIT")
    return jsonify(ok=True)


# ═════════════════════════════════════════════════════════════════════════
#  API:  DOKUMENTE
# ═════════════════════════════════════════════════════════════════════════

def _ext_ok(filename: str) -> bool:
    return "." in filename and \
           filename.rsplit(".", 1)[-1].lower() in ALLOWED_EXT


@app.route("/api/members/<int:member_id>/documents", methods=["POST"])
def api_document_upload(member_id: int):
    with db() as c:
        if not c.execute("SELECT 1 FROM members WHERE id = ?",
                         (member_id,)).fetchone():
            abort(404, "Mitglied nicht gefunden")

    if "file" not in request.files:
        abort(400, "Kein Datei-Upload enthalten")
    f = request.files["file"]
    if not f.filename:
        abort(400, "Kein Dateiname")
    if not _ext_ok(f.filename):
        abort(400, "Dateityp nicht erlaubt")

    original = f.filename
    safe     = secure_filename(original) or "dokument"
    stored   = f"{member_id}_{int(time.time())}_{safe}"
    dest     = DOC_DIR / stored
    f.save(dest)

    doc_type = (request.form.get("doc_type") or "Sonstiges").strip() or "Sonstiges"

    with db() as c:
        c.execute("BEGIN")
        cur = c.execute("""
            INSERT INTO documents (member_id, doc_type, original_name, stored_name, uploaded_at)
            VALUES (?, ?, ?, ?, ?)
        """, (member_id, doc_type, original, stored, now_local()))
        c.execute("COMMIT")
        row = c.execute("SELECT * FROM documents WHERE id = ?",
                        (cur.lastrowid,)).fetchone()
    return jsonify(dict(row)), 201


@app.route("/api/documents/<int:doc_id>/download", methods=["GET"])
def api_document_download(doc_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM documents WHERE id = ?",
                        (doc_id,)).fetchone()
    if not row:
        abort(404, "Dokument nicht gefunden")
    fp = DOC_DIR / row["stored_name"]
    if not fp.exists():
        abort(404, "Datei nicht im Dokumente-Ordner gefunden")
    return send_file(fp, as_attachment=True, download_name=row["original_name"])


@app.route("/api/documents/<int:doc_id>", methods=["DELETE"])
def api_document_delete(doc_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM documents WHERE id = ?",
                        (doc_id,)).fetchone()
        if not row:
            abort(404, "Dokument nicht gefunden")
        c.execute("BEGIN")
        c.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
        c.execute("COMMIT")

    fp = DOC_DIR / row["stored_name"]
    if fp.exists():
        try: fp.unlink()
        except OSError: pass
    return jsonify(ok=True)


# ═════════════════════════════════════════════════════════════════════════
#  API:  EINSTELLUNGEN
# ═════════════════════════════════════════════════════════════════════════

@app.route("/api/settings", methods=["GET"])
def api_settings_get():
    return jsonify(load_settings())


@app.route("/api/settings", methods=["PUT"])
def api_settings_put():
    data = request.get_json(silent=True) or {}
    if not isinstance(data, dict):
        abort(400, "Ungültiges Einstellungs-Format")
    current = load_settings()
    merged  = _deep_merge(current, data)
    save_settings(merged)
    return jsonify(merged)


@app.route("/api/settings/reset", methods=["POST"])
def api_settings_reset():
    save_settings(DEFAULT_SETTINGS)
    return jsonify(DEFAULT_SETTINGS)


# ═════════════════════════════════════════════════════════════════════════
#  OCR-KORREKTUREN  (häufige Verwechslungen in gescannten TuS-PDFs)
# ═════════════════════════════════════════════════════════════════════════

OCR_CITY_MAP = {
    "Lbhne": "Löhne", "Lohne": "Löhne", "Luhne": "Löhne",
    "L6hne": "Löhne", "L0hne": "Löhne", "Loehne": "Löhne",
    "Bunde": "Bünde", "Buende": "Bünde",
    "Lubbecke": "Lübbecke", "Luebbecke": "Lübbecke",
    "Bad Oevnhausen": "Bad Oeynhausen",
    "Bad Oeynbausen": "Bad Oeynhausen",
    "Kirchlengern": "Kirchlengern",
    "Hiddenhausen": "Hiddenhausen",
}

OCR_NAME_MAP = {
    "GéRling":  "Gößling",
    "Bécks":    "Böcks",
    "Luhrmann": "Lührmann",
}


def ocr_correct(text: str) -> str:
    """Wendet bekannte Wort-Korrekturen an (wirkt nur auf ganze Wörter)."""
    if not text:
        return text
    for wrong, right in {**OCR_NAME_MAP, **OCR_CITY_MAP}.items():
        text = re.sub(rf"\b{re.escape(wrong)}\b", right, text)
    return text


# ═════════════════════════════════════════════════════════════════════════
#  FILE-PARSER  (Import: PDF digital + OCR, CSV, TXT, DOCX, XML/SEPA)
# ═════════════════════════════════════════════════════════════════════════

RE_MEMBER_NO_10 = re.compile(r"\b(\d{10})\b")
RE_ZIP_CITY     = re.compile(r"\b(\d{5})\s+([A-ZÄÖÜa-zäöüß][\wäöüß\s\-.]+)")
RE_EMAIL        = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
RE_PHONE        = re.compile(r"(?:\+?\d[\d\s\-/()]{6,})")
RE_DATE_DE      = re.compile(r"\b(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4})\b")
RE_ABTEILUNG    = re.compile(r"Abt\.?\s*\(?\d*\)?\s*([A-Za-zäöüÄÖÜß\-]+)")


def _parse_tus_block(lines: list[str]) -> dict | None:
    """Ein Mitgliederblock aus dem TuS-Textformat.

    Typisches Layout:
        1000000001 Mustermann Musterstraße 12   01.02.1970
                   Max        32584 Löhne       05732 12345
                   max@example.com              05732 98765
    """
    if not lines:
        return None

    m = RE_MEMBER_NO_10.search(lines[0])
    if not m:
        return None
    member_number = m.group(1)

    first, last, street, birth_date = "", "", "", None
    zip_code, city, email = "", "", ""
    phones: list[str] = []

    # Zeile 1: Nummer + Nachname + Straße + Geburtsdatum
    line1 = lines[0].replace(member_number, "", 1).strip()
    dm = RE_DATE_DE.search(line1)
    if dm:
        birth_date = parse_date(dm.group(1))
        line1 = (line1[:dm.start()] + line1[dm.end():]).strip()
    parts1 = line1.split()
    if parts1:
        last = parts1[0]
        street = " ".join(parts1[1:]).strip()

    # Zeile 2: Vorname + PLZ Ort + Telefon
    if len(lines) > 1:
        line2 = lines[1].strip()
        zm = RE_ZIP_CITY.search(line2)
        if zm:
            zip_code = zm.group(1)
            city = zm.group(2).strip()
            before = line2[:zm.start()].strip().split()
            after  = line2[zm.end():].strip()
            if before:
                first = before[0]
            pm = RE_PHONE.search(after)
            if pm:
                phones.append(pm.group(0).strip())
        else:
            parts2 = line2.split()
            if parts2:
                first = parts2[0]

    # Zeile 3+: E-Mail, weitere Telefonnummern
    for extra in lines[2:]:
        em = RE_EMAIL.search(extra)
        if em and not email:
            email = em.group(0).strip()
        for pm in RE_PHONE.finditer(extra):
            p = pm.group(0).strip()
            if p and p not in phones:
                phones.append(p)

    first = ocr_correct(first)
    last  = ocr_correct(last)
    if city:
        city = OCR_CITY_MAP.get(city, city)
    else:
        city = "Löhne"

    if not first or not last:
        return None

    return {
        "member_number": member_number,
        "first_name":    first,
        "last_name":     last,
        "birth_date":    birth_date,
        "street":        street,
        "zip_code":      zip_code or "32584",
        "city":          city,
        "email":         email,
        "phone":         phones[0] if phones else "",
        "gender":        guess_gender(first),
    }


def parse_tus_text(text: str) -> list[dict]:
    """Zerlegt Rohtext in TuS-Blöcke (10-stellige Nr. am Zeilenanfang)."""
    if not text:
        return []
    text = ocr_correct(text)
    lines = [l.rstrip() for l in text.splitlines() if l.strip()]

    members: list[dict] = []
    current: list[str] = []
    for line in lines:
        if RE_MEMBER_NO_10.match(line.strip()):
            if current:
                m = _parse_tus_block(current)
                if m:
                    members.append(m)
            current = [line]
        elif current:
            current.append(line)
    if current:
        m = _parse_tus_block(current)
        if m:
            members.append(m)

    # Duplikate entfernen
    seen: set[str] = set()
    uniq: list[dict] = []
    for m in members:
        k = m.get("member_number") or f"{m['first_name']}|{m['last_name']}"
        if k not in seen:
            seen.add(k)
            uniq.append(m)
    return uniq


def detect_abteilung(text: str) -> str | None:
    if not text:
        return None
    m = RE_ABTEILUNG.search(text)
    return m.group(1).strip() if m else None


def parse_pdf_digital(path: Path) -> str:
    """Text aus digital erzeugten PDFs."""
    if pdfplumber is None:
        return ""
    chunks: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    chunks.append(t)
    except Exception as e:
        app.logger.warning("pdfplumber-Fehler: %s", e)
    return "\n".join(chunks)


def parse_pdf_ocr(path: Path, dpi: int = 400) -> str:
    """OCR für gescannte PDFs — deutsch+englisch, Bildvorverarbeitung."""
    if not OCR_AVAILABLE:
        return ""

    try:
        pages = convert_from_path(str(path), dpi=dpi)
    except Exception as e:
        app.logger.warning("pdf2image-Fehler: %s", e)
        return ""

    out: list[str] = []
    for img in pages:
        # Graustufen → Kontrast ×2 → Schärfe ×2 → Helligkeit +20 % → Binarisierung
        img = img.convert("L")
        img = ImageEnhance.Contrast(img).enhance(2.0)
        img = ImageEnhance.Sharpness(img).enhance(2.0)
        img = ImageEnhance.Brightness(img).enhance(1.2)
        img = img.point(lambda p: 0 if p < 140 else 255, mode="1")

        try:
            text = pytesseract.image_to_string(
                img,
                lang="deu+eng",
                config="--psm 6 -c preserve_interword_spaces=1",
            )
        except Exception as e:
            app.logger.warning("pytesseract-Fehler: %s", e)
            continue

        if text.strip():
            out.append(text)
    return "\n".join(out)


def _decode_bytes(raw: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _map_csv_row(row: dict) -> dict:
    """Spaltennamen auf unser Schema mappen (deutsch + englisch)."""
    def pick(*keys: str) -> str:
        for k in keys:
            for col, val in row.items():
                if col and col.strip().lower() == k.lower() and val:
                    return str(val).strip()
        return ""

    first = pick("vorname", "first_name", "firstname", "given name")
    last  = pick("nachname", "last_name", "lastname", "surname", "familienname")
    if not first and not last:
        full = pick("name", "voller name")
        parts = full.split()
        if len(parts) >= 2:
            first, last = parts[0], " ".join(parts[1:])
        elif len(parts) == 1:
            last = parts[0]

    return {
        "member_number": pick("mitgliedsnr", "mitgliedsnummer", "mitglied", "nummer", "nr"),
        "first_name":    first,
        "last_name":     last,
        "birth_date":    parse_date(pick("geburtsdatum", "geburt", "birth_date", "geb")),
        "gender":        pick("geschlecht", "gender") or guess_gender(first),
        "email":         pick("email", "e-mail", "mail"),
        "phone":         pick("telefon", "phone", "tel", "mobil"),
        "street":        pick("straße", "strasse", "street", "adresse"),
        "zip_code":      pick("plz", "zip", "postleitzahl") or "32584",
        "city":          pick("ort", "stadt", "city") or "Löhne",
        "entry_date":    parse_date(pick("eintritt", "eintrittsdatum", "entry_date")),
    }


def parse_csv(path: Path) -> list[dict]:
    """CSV mit Encoding- und Delimiter-Erkennung."""
    text = _decode_bytes(path.read_bytes())
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        class _D(csv.excel):
            delimiter = ";"
        dialect = _D

    reader = csv.DictReader(io.StringIO(text), dialect=dialect)
    out = []
    for row in reader:
        if not row:
            continue
        mapped = _map_csv_row(row)
        if mapped["first_name"] or mapped["last_name"]:
            out.append(mapped)
    return out


def parse_txt(path: Path) -> list[dict]:
    """TuS-Textformat (mit 10-stelliger Nr.) oder einfach Name-pro-Zeile."""
    text = _decode_bytes(path.read_bytes())
    if RE_MEMBER_NO_10.search(text):
        return parse_tus_text(text)

    members: list[dict] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        parts = line.split(None, 1)
        if len(parts) == 2:
            first, last = parts
        else:
            first, last = "", parts[0]
        members.append({
            "first_name": first,
            "last_name":  last,
            "gender":     guess_gender(first) if first else "",
            "city":       "Löhne",
            "zip_code":   "32584",
        })
    return members


def parse_docx(path: Path) -> list[dict]:
    """Word-Dokumente: Tabellen bevorzugt, Fallback TuS-Text."""
    if python_docx is None:
        return []
    try:
        doc = python_docx.Document(str(path))
    except Exception as e:
        app.logger.warning("python-docx-Fehler: %s", e)
        return []

    members: list[dict] = []
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        header = [c.text.strip().lower() for c in tbl.rows[0].cells]
        for row in tbl.rows[1:]:
            rowd = {header[i]: (row.cells[i].text.strip() if i < len(row.cells) else "")
                    for i in range(len(header))}
            mapped = _map_csv_row(rowd)
            if mapped["first_name"] or mapped["last_name"]:
                members.append(mapped)

    if not members:
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if RE_MEMBER_NO_10.search(text):
            members = parse_tus_text(text)
    return members


def _strip_ns(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _el_text(el) -> str:
    return (el.text or "").strip() if el is not None else ""


def _split_name(full: str) -> tuple[str, str]:
    parts = full.strip().split()
    if not parts:
        return "", ""
    if "," in full:
        a, b = [p.strip() for p in full.split(",", 1)]
        fn = b.split()[0] if b.split() else ""
        return fn, a
    if len(parts) == 1:
        return "", parts[0]
    return parts[0], " ".join(parts[1:])


def parse_xml(path: Path) -> list[dict]:
    """Generische Vereins-XML oder SEPA pain.008."""
    try:
        tree = ET.parse(str(path))
    except ET.ParseError as e:
        app.logger.warning("XML-Parse-Fehler: %s", e)
        return []
    root = tree.getroot()

    members: list[dict] = []

    # SEPA pain.008: DrctDbtTxInf → Dbtr
    for tx in root.iter():
        if _strip_ns(tx.tag) != "DrctDbtTxInf":
            continue
        dbtr = None
        for e in tx.iter():
            if _strip_ns(e.tag) == "Dbtr":
                dbtr = e
                break
        if dbtr is None:
            continue
        name_el = next((c for c in dbtr if _strip_ns(c.tag) == "Nm"), None)
        name = _el_text(name_el)
        pstl = next((c for c in dbtr if _strip_ns(c.tag) == "PstlAdr"), None)
        street = city = zip_ = ""
        if pstl is not None:
            for c in pstl:
                t = _strip_ns(c.tag)
                if t == "StrtNm": street = _el_text(c)
                elif t == "BldgNb": street = (street + " " + _el_text(c)).strip()
                elif t == "TwnNm":  city   = _el_text(c)
                elif t == "PstCd":  zip_   = _el_text(c)
        if name:
            first, last = _split_name(name)
            members.append({
                "first_name": first,
                "last_name":  last,
                "gender":     guess_gender(first),
                "street":     street,
                "city":       city or "Löhne",
                "zip_code":   zip_ or "32584",
            })

    if members:
        return members

    # Generisches Vereins-XML
    for node in root.iter():
        tag = _strip_ns(node.tag)
        if tag not in ("member", "Mitglied", "Mitgliedsdatensatz", "Person"):
            continue
        d = {_strip_ns(ch.tag).lower(): (ch.text or "").strip() for ch in node}
        first = d.get("vorname") or d.get("firstname") or ""
        last  = d.get("nachname") or d.get("lastname") or ""
        if not first and not last:
            continue
        members.append({
            "member_number": d.get("mitgliedsnr") or d.get("id") or "",
            "first_name":    first,
            "last_name":     last,
            "birth_date":    parse_date(d.get("geburtsdatum") or d.get("birth_date")),
            "gender":        d.get("geschlecht") or d.get("gender") or guess_gender(first),
            "email":         d.get("email") or "",
            "phone":         d.get("telefon") or d.get("phone") or "",
            "street":        d.get("strasse") or d.get("straße") or d.get("street") or "",
            "zip_code":      d.get("plz") or "32584",
            "city":          d.get("ort") or d.get("stadt") or "Löhne",
            "entry_date":    parse_date(d.get("eintritt") or d.get("entry_date")),
        })
    return members


# ═════════════════════════════════════════════════════════════════════════
#  API:  IMPORT
# ═════════════════════════════════════════════════════════════════════════

@app.route("/api/import/pdf", methods=["POST"])
def api_import_upload():
    """Akzeptiert PDF/CSV/TXT/DOCX/XML und liefert erkannte Mitglieder als JSON."""
    if "file" not in request.files:
        abort(400, "Keine Datei übergeben")
    f = request.files["file"]
    if not f.filename:
        abort(400, "Kein Dateiname")
    if not _ext_ok(f.filename):
        abort(400, "Dateityp nicht erlaubt")

    ext = f.filename.rsplit(".", 1)[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tf:
        f.save(tf.name)
        tmp = Path(tf.name)

    try:
        abteilung: str | None = None
        members: list[dict]   = []

        if ext == "pdf":
            text = parse_pdf_digital(tmp)
            if (not text or not text.strip()) and OCR_AVAILABLE:
                text = parse_pdf_ocr(tmp)
            abteilung = detect_abteilung(text or "")
            members   = parse_tus_text(text or "")
        elif ext == "csv":
            members = parse_csv(tmp)
        elif ext == "txt":
            members = parse_txt(tmp)
        elif ext in ("docx", "doc"):
            members = parse_docx(tmp)
        elif ext == "xml":
            members = parse_xml(tmp)
        else:
            abort(400, f"Dateityp .{ext} wird nicht unterstützt")
    finally:
        try: tmp.unlink()
        except OSError: pass

    clean: list[dict] = []
    for m in members:
        first = (m.get("first_name") or "").strip()
        last  = (m.get("last_name")  or "").strip()
        if not first and not last:
            continue
        if not m.get("gender") and first:
            m["gender"] = guess_gender(first)
        clean.append(m)

    return jsonify({
        "members_found": clean,
        "count":         len(clean),
        "abteilung":     abteilung,
    })


@app.route("/api/import/confirm", methods=["POST"])
def api_import_confirm():
    """Speichert ausgewählte Mitglieder und ordnet sie optional einer Gruppe zu."""
    data       = request.get_json(silent=True) or {}
    members_in = data.get("members") or []
    group_raw  = data.get("group_id")
    if not isinstance(members_in, list) or not members_in:
        abort(400, "Keine Mitglieder übergeben")

    group_id = int(group_raw) if group_raw and str(group_raw).isdigit() else None

    inserted, skipped = 0, 0
    with db() as c:
        c.execute("BEGIN")
        for m in members_in:
            first = (m.get("first_name") or "").strip()
            last  = (m.get("last_name")  or "").strip()
            if not first or not last:
                skipped += 1
                continue

            number = (m.get("member_number") or "").strip() or next_member_number(c)
            now    = now_local()

            cur = c.execute("""
                INSERT OR IGNORE INTO members (
                    member_number, first_name, last_name, birth_date, gender,
                    email, phone, street, city, zip_code, entry_date, exit_date,
                    color_mark, app_status, notes, is_active, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                number, first, last,
                parse_date(m.get("birth_date")),
                m.get("gender") or guess_gender(first),
                (m.get("email") or "").strip(),
                (m.get("phone") or "").strip(),
                (m.get("street") or "").strip(),
                (m.get("city")   or "Löhne").strip(),
                (m.get("zip_code") or "32584").strip(),
                parse_date(m.get("entry_date")),
                parse_date(m.get("exit_date")),
                "", "Nicht installiert", "",
                0 if m.get("exit_date") else 1,
                now, now,
            ))
            if cur.rowcount:
                inserted += 1
                if group_id:
                    c.execute(
                        "INSERT OR IGNORE INTO member_groups (member_id, group_id) VALUES (?, ?)",
                        (cur.lastrowid, group_id),
                    )
            else:
                skipped += 1
        c.execute("COMMIT")

    return jsonify({"imported": inserted, "skipped": skipped})


# ═════════════════════════════════════════════════════════════════════════
#  API:  DASHBOARD-STATISTIKEN
# ═════════════════════════════════════════════════════════════════════════

AGE_BUCKETS = [
    ("0–6",   0,   6),
    ("7–14",  7,  14),
    ("15–17", 15, 17),
    ("18–30", 18, 30),
    ("31–50", 31, 50),
    ("51–65", 51, 65),
    ("66+",   66, 999),
]

MEMBERSHIP_BUCKETS = [
    ("< 1 Jahr",    0,   0),
    ("1–5 Jahre",   1,   5),
    ("6–10 Jahre",  6,  10),
    ("11–20 Jahre", 11, 20),
    ("20+ Jahre",   21, 999),
]


def _bucket_for(value: float, buckets: list[tuple[str, int, int]]) -> str | None:
    for label, lo, hi in buckets:
        if lo <= value <= hi:
            return label
    return None


@app.route("/api/stats", methods=["GET"])
def api_stats():
    with db() as c:
        total    = c.execute("SELECT COUNT(*) FROM members").fetchone()[0]
        active   = c.execute("SELECT COUNT(*) FROM members WHERE is_active = 1").fetchone()[0]
        inactive = total - active

        gender_rows = c.execute("""
            SELECT COALESCE(NULLIF(gender, ''), 'unbekannt') AS g, COUNT(*) AS n
            FROM members WHERE is_active = 1 GROUP BY g
        """).fetchall()
        gender = {r["g"]: r["n"] for r in gender_rows}

        app_rows = c.execute("""
            SELECT app_status AS s, COUNT(*) AS n
            FROM members WHERE is_active = 1 GROUP BY s
        """).fetchall()
        app_status = {r["s"]: r["n"] for r in app_rows}
        app_users  = app_status.get("Installiert", 0) + app_status.get("Aktiv", 0)

        color_rows = c.execute("""
            SELECT COALESCE(NULLIF(color_mark, ''), 'keine') AS col, COUNT(*) AS n
            FROM members WHERE is_active = 1 GROUP BY col
        """).fetchall()
        colors = {r["col"]: r["n"] for r in color_rows}

        active_rows = c.execute("""
            SELECT id, birth_date, entry_date, email, phone, is_active
            FROM members WHERE is_active = 1
        """).fetchall()

        group_counts = {r["mid"]: r["cnt"] for r in c.execute("""
            SELECT member_id AS mid, COUNT(*) AS cnt
            FROM member_groups GROUP BY member_id
        """).fetchall()}

        age_dist   = {b[0]: 0 for b in AGE_BUCKETS}
        mship_dist = {b[0]: 0 for b in MEMBERSHIP_BUCKETS}
        today = date.today()
        karteileichen = 0

        for r in active_rows:
            if r["birth_date"]:
                try:
                    b = date.fromisoformat(r["birth_date"])
                    age = today.year - b.year - ((today.month, today.day) < (b.month, b.day))
                    lbl = _bucket_for(age, AGE_BUCKETS)
                    if lbl:
                        age_dist[lbl] += 1
                except ValueError:
                    pass
            if r["entry_date"]:
                try:
                    e = date.fromisoformat(r["entry_date"])
                    yrs = (today - e).days / 365.25
                    lbl = _bucket_for(yrs, MEMBERSHIP_BUCKETS)
                    if lbl:
                        mship_dist[lbl] += 1
                except ValueError:
                    pass
            if is_karteileiche(r, group_counts.get(r["id"], 0)):
                karteileichen += 1

        grp_rows = c.execute("""
            SELECT g.id, g.name, g.level,
                (SELECT COUNT(*) FROM member_groups mg WHERE mg.group_id = g.id) AS direct
            FROM groups_tbl g
        """).fetchall()
        groups_dist: dict[str, int] = {}
        has_abteilung = any(r["level"] == 1 for r in grp_rows)
        for r in grp_rows:
            if has_abteilung:
                if r["level"] == 1:
                    groups_dist[r["name"]] = r["direct"]
            else:
                groups_dist[r["name"]] = r["direct"]

        group_total = c.execute("SELECT COUNT(*) FROM groups_tbl").fetchone()[0]

    return jsonify({
        "total":         total,
        "active":        active,
        "inactive":      inactive,
        "karteileichen": karteileichen,
        "groups_count":  group_total,
        "app_users":     app_users,
        "app_status":    app_status,
        "gender":        gender,
        "age":           age_dist,
        "membership":    mship_dist,
        "colors":        colors,
        "groups":        groups_dist,
    })


# ═════════════════════════════════════════════════════════════════════════
#  API:  JAHRESAUSWERTUNG  (+ Jubiläen-Erkennung)
# ═════════════════════════════════════════════════════════════════════════

JUBILEE_YEARS = (5, 10, 15, 20, 25, 30, 40, 50)


def _compute_yearly_stats(year: int) -> dict:
    year_start = date(year, 1, 1).isoformat()
    year_end   = date(year, 12, 31).isoformat()

    with db() as c:
        beginn = c.execute("""
            SELECT COUNT(*) FROM members
            WHERE entry_date IS NOT NULL AND entry_date < ?
              AND (exit_date IS NULL OR exit_date = '' OR exit_date >= ?)
        """, (year_start, year_start)).fetchone()[0]

        zugaenge_rows = c.execute("""
            SELECT * FROM members
            WHERE entry_date IS NOT NULL
              AND entry_date >= ? AND entry_date <= ?
            ORDER BY entry_date, last_name COLLATE NOCASE
        """, (year_start, year_end)).fetchall()

        abgaenge_rows = c.execute("""
            SELECT * FROM members
            WHERE exit_date IS NOT NULL AND exit_date <> ''
              AND exit_date >= ? AND exit_date <= ?
            ORDER BY exit_date, last_name COLLATE NOCASE
        """, (year_start, year_end)).fetchall()

        zugaenge = len(zugaenge_rows)
        abgaenge = len(abgaenge_rows)
        ende     = beginn + zugaenge - abgaenge
        netto    = zugaenge - abgaenge
        netto_pct = (netto / beginn * 100.0) if beginn else 0.0

        monatlich = []
        for month in range(1, 13):
            last_day = date(year, month, monthrange(year, month)[1]).isoformat()
            count = c.execute("""
                SELECT COUNT(*) FROM members
                WHERE entry_date IS NOT NULL AND entry_date <= ?
                  AND (exit_date IS NULL OR exit_date = '' OR exit_date > ?)
            """, (last_day, last_day)).fetchone()[0]
            monatlich.append({"month": month, "count": count})

        # Jubiläen: aktive Mitglieder, deren Eintrittsjahr genau X Jahre zurückliegt
        jubilaeen: list[dict] = []
        jub_rows = c.execute("""
            SELECT id, first_name, last_name, entry_date
            FROM members
            WHERE entry_date IS NOT NULL AND entry_date <> ''
              AND (exit_date IS NULL OR exit_date = '')
        """).fetchall()
        for r in jub_rows:
            try:
                e = date.fromisoformat(r["entry_date"])
            except ValueError:
                continue
            years = year - e.year
            if years in JUBILEE_YEARS:
                jubilaeen.append({
                    "member_id":  r["id"],
                    "first_name": r["first_name"],
                    "last_name":  r["last_name"],
                    "entry_date": r["entry_date"],
                    "years":      years,
                })
        jubilaeen.sort(key=lambda x: (-x["years"], x["last_name"].lower()))

        def _enrich(rows):
            out = []
            for r in rows:
                gs = _load_member_groups(c, r["id"])
                d = dict(r)
                d["groups"] = gs
                out.append(d)
            return out
        zugaenge_list = _enrich(zugaenge_rows)
        abgaenge_list = _enrich(abgaenge_rows)

    return {
        "year":          year,
        "beginn":        beginn,
        "zugaenge":      zugaenge,
        "abgaenge":      abgaenge,
        "ende":          ende,
        "netto":         netto,
        "netto_pct":     round(netto_pct, 1),
        "monatlich":     monatlich,
        "zugaenge_list": zugaenge_list,
        "abgaenge_list": abgaenge_list,
        "jubilaeen":     jubilaeen,
    }


@app.route("/api/stats/yearly", methods=["GET"])
def api_stats_yearly():
    try:
        year = int(request.args.get("year") or date.today().year)
    except ValueError:
        abort(400, "Ungültiges Jahr")
    return jsonify(_compute_yearly_stats(year))


# ═════════════════════════════════════════════════════════════════════════
#  API:  CSV-EXPORT
# ═════════════════════════════════════════════════════════════════════════

@app.route("/api/export/csv", methods=["GET"])
def api_export_csv():
    with db() as c:
        rows = c.execute("""
            SELECT m.member_number, m.first_name, m.last_name, m.birth_date,
                   m.gender, m.email, m.phone, m.street, m.zip_code, m.city,
                   m.entry_date, m.exit_date, m.app_status, m.is_active,
                   m.color_mark, m.notes
            FROM members m
            ORDER BY m.last_name COLLATE NOCASE, m.first_name COLLATE NOCASE
        """).fetchall()

    buf = io.StringIO()
    writer = csv.writer(buf, delimiter=";", quoting=csv.QUOTE_MINIMAL)
    writer.writerow([
        "Mitgliedsnr", "Vorname", "Nachname", "Geburtsdatum", "Geschlecht",
        "E-Mail", "Telefon", "Straße", "PLZ", "Ort",
        "Eintritt", "Austritt", "App-Status", "Aktiv", "Farbe", "Notizen",
    ])
    for r in rows:
        writer.writerow([
            r["member_number"] or "", r["first_name"], r["last_name"],
            r["birth_date"] or "", r["gender"] or "",
            r["email"] or "", r["phone"] or "",
            r["street"] or "", r["zip_code"] or "", r["city"] or "",
            r["entry_date"] or "", r["exit_date"] or "",
            r["app_status"] or "", "ja" if r["is_active"] else "nein",
            r["color_mark"] or "", (r["notes"] or "").replace("\n", " "),
        ])

    out = io.BytesIO(buf.getvalue().encode("utf-8-sig"))
    out.seek(0)
    filename = f"tus-mitglieder_{date.today().isoformat()}.csv"
    return send_file(out, as_attachment=True, download_name=filename,
                     mimetype="text/csv")


# ═════════════════════════════════════════════════════════════════════════
#  API:  JAHRESBERICHT ALS PDF (reportlab)
# ═════════════════════════════════════════════════════════════════════════

def _pdf_person_table(rows: list[dict], primary: str,
                      date_field: str = "entry_date") -> "Table":
    body = ParagraphStyle("tbl", parent=getSampleStyleSheet()["Normal"],
                          fontSize=9, leading=11)
    tbl = [[Paragraph("<b>Name</b>", body),
            Paragraph("<b>Datum</b>", body),
            Paragraph("<b>Gruppen</b>", body)]]
    for r in rows:
        groups = ", ".join(g["name"] for g in (r.get("groups") or []))
        tbl.append([
            Paragraph(f"{r.get('last_name', '')}, {r.get('first_name', '')}", body),
            Paragraph(r.get(date_field) or "", body),
            Paragraph(groups, body),
        ])
    t = Table(tbl, colWidths=[75*mm, 30*mm, 65*mm], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(primary)),
        ("TEXTCOLOR",  (0, 0), (-1, 0), white),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#f5f6fa")]),
        ("BOX",        (0, 0), (-1, -1), 0.3, lightgrey),
        ("LEFTPADDING",  (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


@app.route("/api/export/yearly-pdf", methods=["GET"])
def api_export_yearly_pdf():
    if not REPORTLAB_AVAILABLE:
        abort(500, "reportlab nicht installiert — bitte 'pip install reportlab' ausführen")
    try:
        year = int(request.args.get("year") or date.today().year)
    except ValueError:
        abort(400, "Ungültiges Jahr")

    data     = _compute_yearly_stats(year)
    settings = load_settings()
    primary  = settings.get("colors", {}).get("primary", "#3156a3")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20*mm, bottomMargin=20*mm,
        leftMargin=18*mm, rightMargin=18*mm,
        title=f"Jahresbericht {year} — TuS Gohfeld e.V.",
        author="TuS Gohfeld Mitgliederverwaltung",
    )

    styles = getSampleStyleSheet()
    h1   = ParagraphStyle("h1", parent=styles["Heading1"],
                          textColor=HexColor(primary), fontSize=22, spaceAfter=10)
    h2   = ParagraphStyle("h2", parent=styles["Heading2"],
                          textColor=HexColor(primary), fontSize=15,
                          spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=13)

    story: list[Any] = [
        Paragraph(f"Jahresbericht {year}", h1),
        Paragraph("TuS Gohfeld von 1910 e.V. — Mitgliederentwicklung", body),
        Spacer(1, 6*mm),
    ]

    kpi = [
        [Paragraph(f"<b>{data['beginn']}</b>", body),
         Paragraph(f"<b>{data['zugaenge']}</b>", body),
         Paragraph(f"<b>{data['abgaenge']}</b>", body),
         Paragraph(f"<b>{data['ende']}</b>", body)],
        [Paragraph("Anfangsbestand", body), Paragraph("Zugänge", body),
         Paragraph("Abgänge", body),        Paragraph("Endbestand", body)],
    ]
    kpi_tbl = Table(kpi, colWidths=[42*mm]*4, rowHeights=[15*mm, 8*mm])
    kpi_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(primary)),
        ("TEXTCOLOR",  (0, 0), (-1, 0), white),
        ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE",   (0, 0), (-1, 0), 18),
        ("FONTSIZE",   (0, 1), (-1, 1), 9),
        ("TEXTCOLOR",  (0, 1), (-1, 1), HexColor("#555555")),
        ("BOX",        (0, 0), (-1, -1), 0.5, lightgrey),
        ("INNERGRID",  (0, 0), (-1, -1), 0.3, lightgrey),
    ]))
    story.append(kpi_tbl)
    story.append(Spacer(1, 4*mm))

    arrow = "▲" if data["netto"] >= 0 else "▼"
    story.append(Paragraph(
        f"<b>Nettoveränderung:</b> {arrow} {data['netto']:+d} "
        f"({data['netto_pct']:+.1f} %)",
        body,
    ))

    if data["zugaenge_list"]:
        story.append(Paragraph("Zugänge im Jahr", h2))
        story.append(_pdf_person_table(data["zugaenge_list"], primary, "entry_date"))
    if data["abgaenge_list"]:
        story.append(Paragraph("Abgänge im Jahr", h2))
        story.append(_pdf_person_table(data["abgaenge_list"], primary, "exit_date"))

    if data["jubilaeen"]:
        story.append(PageBreak())
        story.append(Paragraph("Jubiläen", h2))
        tbl = [[Paragraph("<b>Jahre</b>", body),
                Paragraph("<b>Name</b>", body),
                Paragraph("<b>Eintritt</b>", body)]]
        for j in data["jubilaeen"]:
            tbl.append([
                Paragraph(str(j["years"]), body),
                Paragraph(f"{j['last_name']}, {j['first_name']}", body),
                Paragraph(j["entry_date"] or "", body),
            ])
        t = Table(tbl, colWidths=[20*mm, 110*mm, 40*mm], repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor(primary)),
            ("TEXTCOLOR",  (0, 0), (-1, 0), white),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#f5f6fa")]),
            ("BOX",        (0, 0), (-1, -1), 0.3, lightgrey),
        ]))
        story.append(t)

    def _footer(canv, _d):
        canv.saveState()
        canv.setFont("Helvetica", 8)
        canv.setFillColor(HexColor("#777777"))
        canv.drawString(18*mm, 10*mm,
                        f"TuS Gohfeld e.V. — erzeugt am "
                        f"{date.today().strftime('%d.%m.%Y')}")
        canv.drawRightString(A4[0] - 18*mm, 10*mm, f"Seite {_d.page}")
        canv.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    buf.seek(0)
    filename = f"TuS_Jahresbericht_{year}.pdf"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype="application/pdf")


# ═════════════════════════════════════════════════════════════════════════
#  START
# ═════════════════════════════════════════════════════════════════════════

def main() -> None:
    print("━" * 60)
    print("  TuS Gohfeld Mitgliederverwaltung")
    print("━" * 60)
    print(f"  Datenpfad:   {DATA_DIR}")
    print(f"  Datenbank:   {DB_PATH}")
    print(f"  Dokumente:   {DOC_DIR}")
    print(f"  Einstell.:   {SETTINGS}")
    print("━" * 60)
    init_db()
    # host 127.0.0.1 statt 'localhost' (macOS IPv6-Problem)
    app.run(host="127.0.0.1", port=5000, debug=False, use_reloader=False)


if __name__ == "__main__":
    main()
